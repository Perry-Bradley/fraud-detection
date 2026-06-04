"""Report-card computation + PDF rendering.

The gradebook stores raw scores; this module turns them into a per-term report
card: each subject's weighted average (scaled to /20), coefficient-weighted
overall average, class rank, and a letter grade. Pure-Python computation so it's
easy to unit-test; PDF rendering reuses reportlab (already a dependency).
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from decimal import Decimal

from .models import SchoolClass, ClassSubject, Assessment, Grade


def letter_grade(avg20: float | None) -> str:
    if avg20 is None:
        return "-"
    if avg20 >= 18:
        return "A+"
    if avg20 >= 16:
        return "A"
    if avg20 >= 14:
        return "B"
    if avg20 >= 12:
        return "C"
    if avg20 >= 10:
        return "D"
    return "F"


@dataclass
class SubjectResult:
    subject_code: str
    subject_name: str
    coefficient: int
    average20: float | None      # subject average scaled to /20
    weighted: float | None       # average20 * coefficient
    grade: str
    teacher: str = ""


@dataclass
class ReportCard:
    student_id: str
    matricule: str
    full_name: str
    class_name: str
    term: str
    subjects: list[SubjectResult] = field(default_factory=list)
    total_coefficient: int = 0
    total_weighted: float = 0.0
    average20: float | None = None
    grade: str = "-"
    rank: int | None = None
    class_size: int | None = None


def _subject_average20(student, school_class, subject, term) -> float | None:
    """Weighted average of a student's scores in one subject this term, on /20."""
    assessments = Assessment.objects.filter(
        school_class=school_class, subject=subject, term=term
    )
    grades = {
        g.assessment_id: g
        for g in Grade.objects.filter(assessment__in=assessments, student=student)
    }

    num = Decimal("0")
    den = Decimal("0")
    for a in assessments:
        g = grades.get(a.id)
        if g is None or g.score is None or g.is_absent or not a.max_score:
            continue
        pct = (g.score / a.max_score) * a.weight
        num += pct
        den += a.weight
    if den == 0:
        return None
    # mean fraction (0..1) -> scale to /20
    return float(num / den) * 20.0


def compute_report_card(student, term) -> ReportCard:
    try:
        school_class = SchoolClass.objects.get(name=student.class_name)
    except SchoolClass.DoesNotExist:
        school_class = None

    rc = ReportCard(
        student_id=str(student.id),
        matricule=student.matricule,
        full_name=student.full_name,
        class_name=student.class_name,
        term=str(term),
    )
    if school_class is None:
        return rc

    total_coef = 0
    total_weighted = 0.0
    for cs in ClassSubject.objects.select_related("subject", "teacher").filter(
        school_class=school_class
    ):
        avg20 = _subject_average20(student, school_class, cs.subject, term)
        weighted = None if avg20 is None else round(avg20 * cs.coefficient, 2)
        rc.subjects.append(
            SubjectResult(
                subject_code=cs.subject.code,
                subject_name=cs.subject.name,
                coefficient=cs.coefficient,
                average20=None if avg20 is None else round(avg20, 2),
                weighted=weighted,
                grade=letter_grade(avg20),
                teacher=getattr(cs.teacher, "full_name", "") or "",
            )
        )
        if avg20 is not None:
            total_coef += cs.coefficient
            total_weighted += avg20 * cs.coefficient

    rc.total_coefficient = total_coef
    rc.total_weighted = round(total_weighted, 2)
    rc.average20 = round(total_weighted / total_coef, 2) if total_coef else None
    rc.grade = letter_grade(rc.average20)
    return rc


def compute_class_rankings(school_class: SchoolClass, term) -> list[ReportCard]:
    """Report card for every student in the class, ranked by overall average."""
    cards = [compute_report_card(s, term) for s in school_class.students()]
    ranked = sorted(
        cards,
        key=lambda c: (c.average20 if c.average20 is not None else -1),
        reverse=True,
    )
    size = len(ranked)
    for i, c in enumerate(ranked, start=1):
        c.rank = i if c.average20 is not None else None
        c.class_size = size
    return ranked


def report_card_with_rank(student, term) -> ReportCard:
    """A single student's card, but with rank computed against the whole class."""
    try:
        school_class = SchoolClass.objects.get(name=student.class_name)
    except SchoolClass.DoesNotExist:
        return compute_report_card(student, term)
    for card in compute_class_rankings(school_class, term):
        if card.student_id == str(student.id):
            return card
    return compute_report_card(student, term)


def render_report_card_pdf(student, term) -> bytes:
    """Render a printable PDF report card."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    card = report_card_with_rank(student, term)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=18 * mm, bottomMargin=18 * mm,
        leftMargin=16 * mm, rightMargin=16 * mm,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=16)
    small = ParagraphStyle("s", parent=styles["Normal"], fontSize=9, textColor=colors.grey)

    elems = [
        Paragraph("REPORT CARD", title),
        Paragraph("School Management System", small),
        Spacer(1, 8 * mm),
    ]

    meta = [
        ["Student", card.full_name, "Matricule", card.matricule],
        ["Class", card.class_name, "Term", card.term],
    ]
    mt = Table(meta, colWidths=[28 * mm, 55 * mm, 28 * mm, 55 * mm])
    mt.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.grey),
        ("TEXTCOLOR", (2, 0), (2, -1), colors.grey),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elems += [mt, Spacer(1, 6 * mm)]

    rows = [["Subject", "Coef", "Avg /20", "Weighted", "Grade", "Teacher"]]
    for s in card.subjects:
        rows.append([
            s.subject_name,
            str(s.coefficient),
            "-" if s.average20 is None else f"{s.average20:.2f}",
            "-" if s.weighted is None else f"{s.weighted:.2f}",
            s.grade,
            s.teacher or "-",
        ])
    grade_tbl = Table(
        rows, colWidths=[55 * mm, 16 * mm, 22 * mm, 24 * mm, 18 * mm, 35 * mm]
    )
    grade_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d1d5db")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f4f6")]),
        ("ALIGN", (1, 0), (4, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elems += [grade_tbl, Spacer(1, 6 * mm)]

    avg_txt = "-" if card.average20 is None else f"{card.average20:.2f} / 20"
    rank_txt = "-" if card.rank is None else f"{card.rank} of {card.class_size}"
    summary = [
        ["Overall Average", avg_txt, "Grade", card.grade, "Rank", rank_txt],
    ]
    st = Table(summary, colWidths=[34 * mm, 28 * mm, 18 * mm, 16 * mm, 16 * mm, 28 * mm])
    st.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("FONTNAME", (1, 0), (1, 0), "Helvetica-Bold"),
        ("FONTNAME", (3, 0), (3, 0), "Helvetica-Bold"),
        ("FONTNAME", (5, 0), (5, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eff6ff")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#93c5fd")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    elems += [st]

    doc.build(elems)
    return buf.getvalue()


def render_report_card_docx(student, term) -> bytes:
    """Render an EDITABLE Microsoft Word (.docx) report card.

    Unlike the PDF, this opens in Word/Google Docs and every value can be
    edited by hand — useful when a teacher wants to add a comment or correct a
    mark before printing.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    card = report_card_with_rank(student, term)
    doc = Document()

    title = doc.add_heading("REPORT CARD", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("School Management System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student meta as a 2x4 table.
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Light List Accent 1"
    cells = meta.rows[0].cells
    cells[0].text, cells[1].text = "Student", card.full_name
    cells[2].text, cells[3].text = "Matricule", card.matricule
    cells = meta.rows[1].cells
    cells[0].text, cells[1].text = "Class", card.class_name
    cells[2].text, cells[3].text = "Term", card.term

    doc.add_paragraph()

    # Grades table.
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Subject", "Coef", "Avg /20", "Weighted", "Grade", "Teacher"]):
        hdr[i].text = h
    for s in card.subjects:
        row = table.add_row().cells
        row[0].text = s.subject_name
        row[1].text = str(s.coefficient)
        row[2].text = "-" if s.average20 is None else f"{s.average20:.2f}"
        row[3].text = "-" if s.weighted is None else f"{s.weighted:.2f}"
        row[4].text = s.grade
        row[5].text = s.teacher or "-"

    doc.add_paragraph()

    avg_txt = "-" if card.average20 is None else f"{card.average20:.2f} / 20"
    rank_txt = "-" if card.rank is None else f"{card.rank} of {card.class_size}"
    summary = doc.add_paragraph()
    run = summary.add_run(
        f"Overall Average: {avg_txt}      Grade: {card.grade}      Rank: {rank_txt}"
    )
    run.bold = True
    run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
