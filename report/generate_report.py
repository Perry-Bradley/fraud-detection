# -*- coding: utf-8 -*-
"""
CEC418 project report — "AI-Based Fraud Detection in School Fee Payments".

Plain black & white, 12pt body / 14pt sub-heading / 16pt heading, 1.5 spacing,
bullet-driven, targeted at ~20-22 pages with the screenshots embedded.

Drop screenshots into report/images/ using the FIGURES filenames; present
files are embedded, missing ones show a placeholder. Re-run to rebuild.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
try:
    from PIL import Image
except Exception:
    Image = None

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "images")
OUT = os.path.join(HERE, "Fraud_Detection_CEC418_Project_Report.docx")

BLACK = RGBColor(0, 0, 0)
HEADER_FILL = "D9D9D9"   # grayscale — prints fine in B&W
ROW_FILL = "F2F2F2"
FONT = "Times New Roman"
REPO = "https://github.com/Perry-Bradley/fraud-detection"

# ---- figures: (filename, caption) -------------------------------------------
FIGURES = {
    "f_docker": ("docker_desktop.png",
        "The complete platform running under Docker Compose — Postgres, Redis, backend, ML service, Celery worker, frontend, Prometheus, Grafana, Loki, Promtail and pgAdmin."),
    "f_swagger": ("swagger_docs.png",
        "Auto-generated OpenAPI 3.0 (Swagger) documentation produced by drf-spectacular at /api/docs/."),
    "f_fraud_notify": ("fraud_notifications.png",
        "The Fraud Detection console with the live notification drawer surfacing AI-flagged suspicious fee payments."),
    "f_fraud_trend": ("fraud_trend.png",
        "Fraud Detection analytics — the 30-day detection trend and the most common anomaly reasons."),
    "f_circleci": ("ci_circleci.png",
        "CI/CD pipeline execution on CircleCI for the fraud-detection repository — the ml-tests, frontend-build, backend-tests and deploy jobs."),
    "f_grafana": ("grafana_dashboard.png",
        "Grafana — the auto-provisioned “SMS Overview” dashboard showing live backend request-rate and p95 latency sourced from Prometheus."),
    "f_prometheus": ("prometheus_targets.png",
        "Prometheus — the targets page confirming the Django backend, the ML service and Prometheus itself are all being scraped and reporting UP."),
    "f_railway": ("railway_production.png",
        "The live production deployment on Railway — the backend, ML fraud-detection and frontend services, plus managed PostgreSQL and Redis, all online and deployed from the GitHub repository."),
    "f_dash_kpi": ("dashboard_kpis.png",
        "School-management dashboard KPIs — collection rate, payment methods, active students and the anomaly count flagged by the AI."),
    "f_dash_defaulters": ("dashboard_defaulters.png",
        "Dashboard — top fee defaulters and the real-time activity feed, including events flagged by the AI detector."),
}
_fig = {"n": 0}


# ---- helpers -----------------------------------------------------------------
def set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_borders(cell, color="808080", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(a); r.append(it); r.append(s); r.append(e)
    return run


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "000000"); rPr.append(col)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT); rPr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; run.append(t)
    link.append(run); paragraph._p.append(link)
    return link


# ---- document ---------------------------------------------------------------
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
    st = doc.styles[name]
    st.font.name = FONT; st.font.size = Pt(size); st.font.color.rgb = BLACK; st.font.bold = True
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.line_spacing = 1.5

lb = doc.styles["List Bullet"]
lb.font.name = FONT; lb.font.size = Pt(12); lb.font.color.rgb = BLACK
lb.paragraph_format.line_spacing = 1.5; lb.paragraph_format.space_after = Pt(3)

sec = doc.sections[0]
sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.1); sec.right_margin = Inches(1.1)


def body(text, italic=False, size=None, align=None, after=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.color.rgb = BLACK
    if size: r.font.size = Pt(size)
    if align is not None: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead); r.bold = True; r.font.color.rgb = BLACK
    r2 = p.add_run(text); r2.font.color.rgb = BLACK
    return p


def h1(text): return doc.add_heading(text, level=1)
def h2(text): return doc.add_heading(text, level=2)
def pbreak(): doc.add_page_break()


def add_table(headers, rows, widths=None, fs=11):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        p = hdr[i].paragraphs[0]; p.paragraph_format.line_spacing = 1.0
        run = p.add_run(htext); run.bold = True; run.font.size = Pt(fs)
        run.font.name = FONT; run.font.color.rgb = BLACK
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val)); run.font.size = Pt(fs); run.font.name = FONT; run.font.color.rgb = BLACK
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_figure(key, max_width_in=5.7, max_height_in=3.5):
    _fig["n"] += 1
    fname, caption = FIGURES[key]
    path = os.path.join(IMG_DIR, fname)
    if os.path.exists(path):
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # cap by width, but if the image is tall, cap by height instead so
            # portrait screenshots do not overflow the page.
            w_in = max_width_in
            if Image is not None:
                with Image.open(path) as im:
                    iw, ih = im.size
                if (ih / iw) * max_width_in > max_height_in:
                    p.add_run().add_picture(path, height=Inches(max_height_in))
                else:
                    p.add_run().add_picture(path, width=Inches(w_in))
            else:
                p.add_run().add_picture(path, width=Inches(w_in))
        except Exception:
            _placeholder(fname)
    else:
        _placeholder(fname)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10); cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.line_spacing = 1.0
    try: cap.style = doc.styles["Caption"]
    except KeyError: pass
    r = cap.add_run("Figure "); r.bold = True
    add_field(cap, r" SEQ Figure \* ARABIC ")
    r2 = cap.add_run(": "); r2.bold = True
    r3 = cap.add_run(caption); r3.italic = True
    for run in cap.runs:
        run.font.size = Pt(10.5); run.font.name = FONT; run.font.color.rgb = BLACK


def _placeholder(fname):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; cell.width = Inches(5.7)
    set_cell_borders(cell)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(22)
    r = p.add_run("IMAGE PLACEHOLDER\n"); r.bold = True; r.font.size = Pt(12); r.font.name = FONT; r.font.color.rgb = BLACK
    r2 = p.add_run(f"Save the screenshot as  report/images/{fname}  then re-run generate_report.py")
    r2.font.size = Pt(10); r2.italic = True; r2.font.name = FONT; r2.font.color.rgb = BLACK


# =============================================================================
# COVER
# =============================================================================
def cover_line(text, size, bold=False, italic=False, after=4, before=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = FONT; r.font.color.rgb = BLACK
    return p

cover_line("REPUBLIC OF CAMEROON  ·  Peace – Work – Fatherland", 11, italic=True, before=6, after=14)
cover_line("UNIVERSITY OF BUEA", 18, bold=True)
cover_line("COLLEGE OF TECHNOLOGY (COT)", 14, bold=True)
cover_line("DEPARTMENT OF COMPUTER ENGINEERING", 12, bold=True, after=18)

_p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
crest = os.path.join(IMG_DIR, "ub_logo.png")
if os.path.exists(crest):
    _p.add_run().add_picture(crest, width=Inches(1.2))
else:
    cr = _p.add_run("[ University crest ]"); cr.italic = True; cr.font.size = Pt(10); cr.font.name = FONT
_p.paragraph_format.space_after = Pt(18)

cover_line("CEC418 — SOFTWARE CONSTRUCTION AND EVOLUTION", 13, bold=True, after=14)

for line in ("AI-BASED FRAUD DETECTION IN", "SCHOOL FEE PAYMENTS"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line); r.font.size = Pt(19); r.bold = True; r.font.name = FONT; r.font.color.rgb = BLACK
cover_line("A machine-learning module built within a School Management System", 12, italic=True, after=4)
cover_line("Project Report", 12, italic=True, after=24)

tbl = doc.add_table(rows=0, cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
def kv(label, value):
    cells = tbl.add_row().cells
    pl = cells[0].paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rl = pl.add_run(label + "   "); rl.bold = True; rl.font.size = Pt(12); rl.font.name = FONT; rl.font.color.rgb = BLACK
    rr = cells[1].paragraphs[0].add_run(value); rr.font.size = Pt(12); rr.font.name = FONT; rr.font.color.rgb = BLACK
kv("Submitted by:", "Sepo Perry-Bradley Dinga")
kv("Matricule:", "CT23A145  ·  BTech 400")
kv("Course Instructor:", "Mr. Kometa Denis")
kv("Academic Year:", "2025 / 2026")
kv("Date:", "June 2026")
for row in tbl.rows:
    row.cells[0].width = Inches(2.1); row.cells[1].width = Inches(3.3)

pbreak()

# =============================================================================
# ABSTRACT + TOC
# =============================================================================
h1("Abstract")
body("School fees represent a large flow of cash through secondary schools, yet collection is often manual "
     "and weakly controlled, leaving room for errors and internal fraud. This project delivers an AI-based "
     "fraud-detection capability for school-fee payments, built as a module inside a wider School "
     "Management System (SMS). Every payment — whether recorded by a bursar or made online by a "
     "parent — is scored in real time by a machine-learning service (an Isolation Forest) that flags "
     "suspicious transactions for staff review. The platform is engineered as three small services (a "
     "Django REST backend, a React single-page application and a FastAPI machine-learning microservice) "
     "backed by PostgreSQL and Redis, and is wrapped in a complete DevOps toolchain. This report explains "
     "how the project demonstrates the CEC418 “design → construct → test” cycle, with "
     "emphasis on the fraud-detection machine-learning aspect and the DevOps tooling that builds, ships and "
     "observes it.")

h1("Table of Contents")
p = doc.add_paragraph(); add_field(p, 'TOC \\o "1-2" \\h \\z \\u')
h1("List of Figures")
p = doc.add_paragraph(); add_field(p, 'TOC \\h \\z \\c "Figure"')
body("(In Word, press Ctrl+A then F9 to fill in the contents, the figure list and the page numbers.)",
     italic=True, size=10.5)

pbreak()

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
h1("1. Introduction")

h2("1.1 Background")
bullet(" each term, schools collect substantial fees, frequently in cash and recorded by hand in ledgers "
       "or spreadsheets.", lead="Manual collection:")
bullet(" balances are hard to reconcile, receipts can be lost or forged, and there is no automatic way to "
       "notice a suspicious payment.", lead="Weak controls:")
bullet(" because the bursary handles cash directly, fee collection is a natural point for internal financial "
       "fraud — duplicate receipts, altered amounts, or unusual payment patterns.", lead="Fraud risk:")
bullet(" the project therefore focuses on detecting fraud during school-fee payment, delivered as a module "
       "within a broader School Management System.", lead="This project:")

h2("1.2 Problem Statement")
body("The core problem is to automatically and reliably detect potentially fraudulent or anomalous "
     "fee payments at the moment they are recorded, without slowing down day-to-day collection, and to "
     "present those detections to staff in a way they can act on.")

h2("1.3 Aim and Objectives")
body("The aim is to construct a secure, observable fee-payment system whose distinguishing feature is "
     "machine-learning fraud detection. The specific objectives are to:")
bullet(" record fee payments and automatically generate unique, downloadable PDF receipts;")
bullet(" compute real-time student balances and outstanding amounts;")
bullet(" accept online mobile-money payments (via the Campay gateway, in sandbox/stub mode);")
bullet(" score every payment with a machine-learning anomaly detector and surface flagged payments through "
       "a review workflow;")
bullet(" provide staff dashboards, reporting and an immutable audit trail; and")
bullet(" demonstrate the full CEC418 “design → construct → test” cycle and an "
       "industry-standard DevOps toolchain.")

h2("1.4 Scope")
bullet(" the software (backend, frontend, machine-learning service and database), its fraud-detection logic, "
       "and the complete DevOps tooling that builds and runs it.", lead="In scope:")
bullet(" physical hardware procurement and a fully settled production payment integration; the Campay "
       "gateway runs in sandbox/stub mode for demonstration.", lead="Out of scope:")

pbreak()

# =============================================================================
# 2. RELATION TO CEC418
# =============================================================================
h1("2. How the Project Relates to CEC418")
body("CEC418 (Software Construction and Evolution) teaches construction around four fundamentals — "
     "minimising complexity, anticipating change, constructing for verification, and reuse & standards "
     "— together with software evolution, managing construction, and construction tools. The project "
     "was deliberately built to exercise each of these; Table 1 maps the course concepts to concrete "
     "features.")
add_table(
    ["CEC418 concept", "How the project demonstrates it"],
    [
        ("Minimising complexity", "Three small single-responsibility services; the ML model isolated in its own service; thin views over fat models; shared permission classes."),
        ("Anticipating change", "Twelve-factor configuration via environment variables; per-environment Helm values; images tagged per commit."),
        ("Constructing for verification", "Unit tests, an immutable audit log, real-time ML anomaly scoring, structured JSON logging, CodeQL and Trivy scanning."),
        ("Reuse & standards", "OpenAPI schema (drf-spectacular), a reusable Helm chart, conventional REST and 12-factor configuration."),
        ("Software evolution", "Dependabot dependency PRs, a full test re-run on every build, GitOps continuous delivery via Argo CD."),
        ("Managing construction", "A defined design→construct→test lifecycle; IaC provisioning (Terraform) and configuration management (Ansible)."),
        ("Construction tools", "Docker, Kubernetes, Helm, Argo CD, Prometheus/Grafana/Loki, CI/CD — detailed in Section 5."),
    ],
    widths=[1.9, 4.0],
)
body("The “design → construct → test” cycle was applied iteratively rather than as a "
     "single waterfall pass. Each capability went through its own small loop:")
bullet(" identify the actors, data and rules (for fraud detection: what makes a payment “unusual”).", lead="Design:")
bullet(" implement the feature backend-first, expose it through the REST API, then consume it in the UI.", lead="Construct:")
bullet(" verify with automated tests and validate against the running system and the audit trail.", lead="Test:")

pbreak()

# =============================================================================
# 3. SYSTEM OVERVIEW
# =============================================================================
h1("3. System Overview and Architecture")

h2("3.1 Architecture")
body("The platform is composed of three application services and two data services behind an Nginx reverse "
     "proxy:")
bullet(" Django + DRF on port 5000 — authentication, students, fees, payments, receipts, audit, "
       "reporting, and the bridge to the ML service.", lead="Backend: ")
bullet(" FastAPI on port 8000 — the fraud-detection model, exposed as a stateless scoring endpoint.",
       lead="ML service: ")
bullet(" React + Vite on port 3000 — a staff console and a student/parent portal behind one login.",
       lead="Frontend: ")
bullet(" PostgreSQL stores all records; Redis backs caching and the Celery background-task queue.",
       lead="Data tier: ")
body("Although the host application is a full School Management System (with students, gradebook, "
     "attendance, exams, timetabling and admissions modules), this report concentrates on the fee-payment "
     "and fraud-detection core. The entire stack runs locally with a single command (make up) through "
     "Docker Compose (Figure 1).")
add_figure("f_docker")

h2("3.2 Technology Stack")
add_table(
    ["Layer", "Technology", "Version"],
    [
        ("Frontend", "React + Vite + Recharts + Axios", "React 18.3 / Vite 5.2"),
        ("Backend", "Django + DRF + SimpleJWT", "Django 5.0.6 / DRF 3.15"),
        ("Background jobs", "Celery + Redis", "Celery 5.4 / Redis 7"),
        ("ML service", "FastAPI + scikit-learn (Isolation Forest)", "FastAPI 0.111 / sklearn 1.5"),
        ("Database", "PostgreSQL", "15-alpine"),
        ("API documentation", "drf-spectacular (OpenAPI 3 / Swagger)", "0.27.2"),
    ],
    widths=[1.4, 3.1, 1.7],
)
body("The REST API is self-documenting: drf-spectacular generates an OpenAPI 3.0 schema served as "
     "interactive Swagger documentation (Figure 2), an example of building to a recognised standard.")
add_figure("f_swagger")

pbreak()

# =============================================================================
# 4. FRAUD DETECTION (ML)  -- focus
# =============================================================================
h1("4. Fraud Detection: The Machine-Learning Core")
body("Automated fraud detection is the distinguishing feature of the project. Every payment is scored as "
     "it is recorded by a dedicated FastAPI microservice exposing POST /detect-anomaly, which returns "
     "{ is_anomalous, score, reason }. Isolating the model in its own service keeps the heavy scientific "
     "stack (scikit-learn, NumPy, Pandas) out of the web backend and lets it scale and fail independently.")

h2("4.1 Why a Separate Machine-Learning Service")
bullet(" the ML runtime and its large dependencies do not bloat or destabilise the web backend.",
       lead="Separation of concerns: ")
bullet(" the model can be scaled, restarted or replaced without touching the application.", lead="Independent lifecycle: ")
bullet(" a well-defined HTTP contract between backend and model is reusable and testable on its own.",
       lead="Clear interface: ")

h2("4.2 The Detection Model")
body("Detection is a hybrid model that adapts to how much payment history a student has:")
bullet(" if a student has fewer than ten prior payments, a rule-based heuristic is used. It flags a "
       "same-day duplicate amount (high confidence) or an amount whose z-score exceeds three, and it never "
       "flags a first-ever payment.", lead="Cold-start heuristic: ")
bullet(" with enough history, a scikit-learn IsolationForest (100 trees, contamination set to "
       "“auto”, and a fixed random seed for reproducibility) is trained leave-one-out on the "
       "student’s history and used to score the new payment.", lead="Isolation Forest: ")
body("Each payment is represented by a small, interpretable feature vector:")
add_table(
    ["Feature", "Meaning"],
    [
        ("amount", "The raw payment amount."),
        ("amount / mean", "How large the payment is relative to the student’s average."),
        ("z-score", "How many standard deviations the amount is from the student’s mean."),
        ("history length", "How many prior payments the student has made."),
        ("avg gap (days)", "The average age/spacing of the student’s past payments."),
    ],
    widths=[1.8, 4.1],
)

h2("4.3 Scoring and Reasons")
bullet(" the model produces a raw anomaly score that is normalised to the range 0–1; higher means more "
       "suspicious.", lead="Score: ")
bullet(" a confirmed-anomalous payment is reported with a floor score so genuine flags are not lost in "
       "rounding.", lead="Threshold: ")
bullet(" each flag carries a human-readable explanation — for example “Duplicate amount "
       "detected” or “Unusual amount or frequency” — so staff understand why it was "
       "raised.", lead="Reason: ")
body("In practice the most common reasons observed are same-day duplicate transactions, rapid successive "
     "payments, amounts well outside the class average, and round-number anomalies.")

h2("4.4 Integration with Payments")
bullet(" the backend bridge score_payment() assembles a request from the student’s last 50 payments "
       "and POSTs it to the ML service.", lead="On every payment: ")
bullet(" the returned is_anomalous, score and reason are stored directly on the Payment record alongside a "
       "review status (open / investigating / confirmed / dismissed).", lead="Persisted result: ")
bullet(" the same scoring path is reused by in-person payments, the online mobile-money webhook, an admin "
       "“simulate” action and a periodic reconciliation task — a clear case of construction "
       "reuse, ensuring identical audit logs and notifications everywhere.", lead="One path, many sources: ")

h2("4.5 Resilience (Fail-Open Design)")
bullet(" the call to the ML service uses a strict three-second timeout.", lead="Timeout: ")
bullet(" if the model is slow or down, the error is logged and the payment is treated as not anomalous, so "
       "a cashier is never blocked from recording a payment.", lead="Fail-open: ")
bullet(" this deliberate error-handling and fault-tolerance choice reflects the construction-technologies "
       "unit of CEC418 — availability of the core service is prioritised over a single non-critical "
       "check.", lead="Course link: ")

h2("4.6 Detection in Action")
body("When a payment is flagged, the system raises a real-time notification to administrators and lists the "
     "payment on the Fraud Detection console, where staff triage it through the review workflow (Figure 3). "
     "The console also summarises detection trends and ranks the most common anomaly reasons (Figure 4).")
add_figure("f_fraud_notify")
add_figure("f_fraud_trend")

pbreak()

# =============================================================================
# 5. DEVOPS
# =============================================================================
h1("5. The DevOps Toolchain")
body("A central goal was to wrap the application in a complete DevOps toolchain that carries a code change "
     "from a commit all the way to a running, observed service — automatically. Table 3 indexes every "
     "tool; the subsections that follow group them by stage.")
add_table(
    ["Stage", "Tool", "Purpose"],
    [
        ("Source / VCS", "Git + GitHub", "Version control; single source of truth for CI and GitOps."),
        ("CI (active)", "CircleCI", "Run backend, ML and frontend tests on every push."),
        ("CI / CD (mirror)", "GitHub Actions + GHCR", "Test, build, scan and publish images to the registry."),
        ("Dependency updates", "Dependabot", "Weekly automated dependency-bump pull requests."),
        ("Build / package", "Docker", "Multi-stage, non-root container images per service."),
        ("Security scanning", "Trivy + CodeQL", "Image CVE scanning and semantic static analysis (SAST)."),
        ("Infrastructure as Code", "Terraform", "Provision the cluster, managed database and registry."),
        ("Configuration mgmt", "Ansible", "Install tooling and deploy the application."),
        ("Orchestration", "Kubernetes", "Run the containers in the cloud."),
        ("Packaging (k8s)", "Helm", "Templated, per-environment deployments."),
        ("GitOps CD", "Argo CD", "Continuously sync the cluster to Git."),
        ("Secrets", "HashiCorp Vault", "Runtime secret storage and injection."),
        ("Observability", "Prometheus + Grafana + Loki", "Metrics, dashboards and aggregated logs."),
    ],
    widths=[1.6, 1.9, 2.4],
)

h2("5.1 Version Control and Containerization")
bullet(" the whole project lives in one Git repository on GitHub, the single source of truth that both CI "
       "and the GitOps controller observe.", lead="Git + GitHub: ")
bullet(" each service ships a Dockerfile; the backend and frontend use multi-stage builds, all images run "
       "as a non-root user, and every service declares a container HEALTHCHECK.", lead="Docker: ")
bullet(" docker-compose.yml brings the whole platform up with one command, on a private network with "
       "health-gated start-up ordering (Figure 1).", lead="Docker Compose: ")

h2("5.2 Continuous Integration and Delivery")
bullet(" the course required demonstrating tools beyond GitHub’s own ecosystem, so CircleCI was "
       "chosen as the active CI provider. This also keeps the pipeline independent of the source host and "
       "avoids vendor lock-in.", lead="Why CircleCI: ")
bullet(" the CircleCI pipeline runs backend-tests (Python 3.11 with Postgres and Redis service containers, "
       "running migrations and the Django test suite), ml-tests (pytest) and frontend-build (a Vite "
       "production build) on every push (Figure 5).", lead="What it runs: ")
bullet(" the repository also retains GitHub Actions workflows as a mirror that builds the Docker images, "
       "scans them with Trivy, and publishes them to the GitHub Container Registry tagged with the commit "
       "SHA for full traceability — so the delivery path is portable across both providers.",
       lead="GitHub Actions (mirror): ")
add_figure("f_circleci")

h2("5.3 Security and Supply-Chain")
bullet(" scans every built image for CRITICAL and HIGH CVEs; results are uploaded as SARIF to the GitHub "
       "Security tab.", lead="Trivy: ")
bullet(" provides semantic static application security testing (SAST) of the Python and JavaScript code.",
       lead="CodeQL: ")
bullet(" opens weekly pull requests to patch pip, npm, Docker base-image and GitHub-Actions dependencies "
       "across all three services.", lead="Dependabot: ")

h2("5.4 Infrastructure as Code and Orchestration")
bullet(" provisions the cloud footprint on DigitalOcean declaratively — a managed Kubernetes cluster "
       "(autoscaling), a managed PostgreSQL cluster firewalled to the cluster, a private registry, and "
       "in-cluster add-ons — all created with terraform apply and removed with terraform destroy.",
       lead="Terraform: ")
bullet(" then configures and deploys onto the cluster: it installs kubectl and Helm, creates the secrets, "
       "registers the Argo CD application, and verifies the rollout.", lead="Ansible: ")
bullet(" runs the backend at two replicas behind a HorizontalPodAutoscaler (2–8 pods at 70% CPU), "
       "Postgres as a StatefulSet with a persistent volume, a single-replica Celery beat to avoid duplicate "
       "scheduling, an Ingress providing Let’s Encrypt TLS, and a NetworkPolicy that restricts "
       "database access to backend pods only.", lead="Kubernetes: ")

h2("5.5 Packaging and GitOps Delivery")
bullet(" the same workloads are packaged as a reusable chart with per-environment values files "
       "(production and development), so a new environment is just a data change.", lead="Helm: ")
bullet(" watches the chart in Git on the main branch and keeps the cluster automatically synchronised with "
       "prune and self-heal enabled — every push to main is delivered with no manual step, and any "
       "manual drift is reverted to match Git.", lead="Argo CD: ")

h2("5.6 Secrets and Observability")
bullet(" stores secrets and injects them into backend pods at runtime, rather than baking them into images "
       "or manifests.", lead="HashiCorp Vault: ")
bullet(" scrapes the backend and ML service every fifteen seconds.", lead="Prometheus: ")
bullet(" ship every container’s logs to a central store with a retention window.", lead="Promtail + Loki: ")
bullet(" is auto-provisioned with both data sources and an overview dashboard, giving one pane for metrics "
       "and logs — directly supporting the CEC418 emphasis on debugging and verification.",
       lead="Grafana: ")

h2("5.7 Production Deployment (Railway)")
body("The repository includes a full Kubernetes / Terraform / Argo CD path (Sections 5.4–5.5) that "
     "demonstrates the orchestration toolchain. The live production system itself is deployed on Railway, "
     "a platform-as-a-service that builds each service straight from the GitHub repository (each service "
     "points at a different root directory of the monorepo):")
bullet(" the backend (Django API), the ML fraud-detection API and the React frontend are deployed from the "
       "same repository, alongside Railway-managed PostgreSQL and Redis.", lead="Services: ")
bullet(" Railway’s GitHub integration redeploys every service on each push to main, while CircleCI "
       "gates quality — a lightweight continuous-delivery pipeline.", lead="Auto-deploy: ")
bullet(" the frontend and backend expose public HTTPS domains; the ML service stays private and is reached "
       "internally by the backend over the private network.", lead="Networking: ")
bullet(" Figure 6 shows the production environment with every service reporting healthy.", lead="All online: ")
add_figure("f_railway")

h2("5.8 The End-to-End Flow")
body("Putting it together, a single push to main flows automatically through the whole chain:")
bullet(" a developer pushes to GitHub;")
bullet(" CircleCI (and, when enabled, GitHub Actions) runs the full test suite and scans the images;")
bullet(" passing images are published to the registry, tagged by commit;")
bullet(" Argo CD notices the updated Helm chart in Git and syncs it onto the Terraform-provisioned, "
       "Ansible-configured Kubernetes cluster;")
bullet(" the running services emit metrics and logs to Prometheus, Loki and Grafana, with secrets supplied "
       "by Vault.")

pbreak()

# =============================================================================
# 6. VERIFICATION & RESULTS
# =============================================================================
h1("6. Verification, Results and Evolution")

h2("6.1 Verification")
bullet(" Django test cases verify, for example, that receipt numbers are auto-generated and unique.",
       lead="Backend tests: ")
bullet(" pytest tests assert that a first payment is never flagged, that a same-day duplicate is flagged "
       "with a high score, and that a consistent history is not flagged.", lead="ML tests: ")
bullet(" the Vite production build runs in CI as a compile-time smoke test of the whole interface.",
       lead="Frontend build: ")
bullet(" CodeQL, Trivy and Dependabot add continuous, automated verification of code quality and "
       "supply-chain health.", lead="Continuous checks: ")

h2("6.2 Results")
body("The constructed system meets every objective from Section 1.3. Payments are recorded with unique PDF "
     "receipts; balances and outstanding amounts are computed live; online mobile-money payments work in "
     "sandbox mode; and every payment is scored for fraud with flagged cases surfaced for review. The "
     "management dashboards summarise collection performance and the anomaly count detected by the AI "
     "(Figures 7 and 8).")
add_figure("f_dash_kpi")
add_figure("f_dash_defaulters")

h2("6.3 Evolution and Maintainability")
bullet(" the system began as a fee tracker and grew to include gradebook, attendance, exams, timetabling "
       "and admissions modules without destabilising the payments-and-fraud core — perfective and "
       "adaptive evolution.", lead="Designed to grow: ")
bullet(" Dependabot PRs, the security scans and the full CI test re-run form a continuous preventive-"
       "maintenance loop.", lead="Preventive maintenance: ")
bullet(" because images are tagged by commit and Argo CD records the deployed revision, any behaviour can "
       "be traced to its source change and rolled back — satisfying the change-impact and traceability "
       "goals of the evolution unit.", lead="Traceability: ")

pbreak()

# =============================================================================
# 7. CONCLUSION
# =============================================================================
h1("7. Conclusion and Future Work")
body("The project successfully demonstrates the CEC418 “design → construct → test” "
     "cycle on a realistic application whose distinguishing feature is machine-learning fraud detection for "
     "school-fee payments. It scores every payment in real time, surfaces suspicious cases through a review "
     "workflow, and is wrapped in a complete DevOps toolchain that embodies the course principles of "
     "minimising complexity, anticipating change, constructing for verification, and reuse and standards.")
body("Future work includes:")
bullet(" integrating a fully settled production payment provider in place of the sandbox gateway;")
bullet(" retraining the anomaly model on real historical data and adding periodic model evaluation;")
bullet(" raising automated-test coverage, including frontend unit and end-to-end tests; and")
bullet(" adding Prometheus alerting on top of the existing dashboards.")

pbreak()

# =============================================================================
# REFERENCES
# =============================================================================
h1("References")

# project links first (with clickable hyperlinks)
def ref_link(idx, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"[{idx}]  ").bold = True
    p.add_run(label + " ")
    add_hyperlink(p, url, url)

def ref_text(idx, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"[{idx}]  ").bold = True
    p.add_run(text)

ref_link(1, "Project source-code repository (GitHub):", REPO)
ref_link(2, "Continuous-integration pipelines (GitHub Actions):", REPO + "/actions")
ref_link(3, "Issue tracker and pull requests:", REPO + "/pulls")
ref_text(4, "R. E. Fairley, Guide to the Software Engineering Body of Knowledge (SWEBOK), Version 3.0. "
            "Piscataway, NJ: IEEE, 2014.")
ref_text(5, "B. Meyer, Object-Oriented Software Construction, 3rd ed. ISE Inc., 2004.")
ref_text(6, "X. Peng, Software Construction. Software School, Fudan University, China, 2009.")
ref_text(7, "F. T. Liu, K. M. Ting and Z.-H. Zhou, “Isolation Forest,” in Proc. IEEE ICDM, 2008, "
            "pp. 413–422.")
ref_link(8, "scikit-learn, “IsolationForest” documentation:",
         "https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html")
ref_link(9, "Django and Django REST Framework documentation:", "https://www.django-rest-framework.org/")
ref_link(10, "FastAPI documentation:", "https://fastapi.tiangolo.com/")
ref_link(11, "Docker documentation:", "https://docs.docker.com/")
ref_link(12, "Kubernetes documentation:", "https://kubernetes.io/docs/")
ref_link(13, "Helm documentation:", "https://helm.sh/docs/")
ref_link(14, "Terraform documentation:", "https://developer.hashicorp.com/terraform/docs")
ref_link(15, "Argo CD documentation:", "https://argo-cd.readthedocs.io/")
ref_link(16, "Prometheus, Grafana and Loki documentation:", "https://grafana.com/docs/")

# =============================================================================
# FOOTER — page number only, no header, hidden on the cover page
# =============================================================================
section = doc.sections[0]
section.different_first_page_header_footer = True
fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_field(fp, "PAGE")
for r in fp.runs:
    r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = BLACK

doc.save(OUT)
print("Saved:", OUT)
print("Figures embedded:", sum(1 for k,(f,_) in FIGURES.items() if os.path.exists(os.path.join(IMG_DIR,f))),
      "/", len(FIGURES))
